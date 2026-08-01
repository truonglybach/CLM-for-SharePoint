"""Shared document-to-text extraction with local OCR fallback."""
from __future__ import annotations

import io
import logging
from pathlib import Path

log = logging.getLogger(__name__)
_MIN_TEXT_CHARS = 20
_OCR_DPI = 300
_IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp")

def extract_text(filename: str, data: bytes, *, ocr_fallback: bool = True) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        text = _pdf_text_layer(data)
        if len(text.strip()) >= _MIN_TEXT_CHARS or not ocr_fallback:
            return text
        log.info("PDF text layer empty/short; falling back to OCR: %s", filename)
        return _ocr_pdf(data)
    if suffix == ".docx":
        return _docx_text(data)
    if suffix in _IMAGE_EXTS:
        return _ocr_image(data) if ocr_fallback else ""
    raise ValueError(f"Unsupported file type for text extraction: {filename}")

def _pdf_text_layer(data: bytes) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("Install PyMuPDF (pip install pymupdf) for PDF support.") from exc
    with fitz.open(stream=data, filetype="pdf") as doc:
        return "\n".join(page.get_text() for page in doc)

def _docx_text(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("Install python-docx for DOCX support.") from exc
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)

def _ocr_engine():
    try:
        import pytesseract
        from PIL import Image  # noqa
    except ImportError as exc:
        raise RuntimeError("OCR fallback needs pytesseract + Pillow.") from exc
    try:
        pytesseract.get_tesseract_version()
    except Exception as exc:
        raise RuntimeError("Tesseract binary not found. Install it locally; do not route to hosted OCR API.") from exc
    return pytesseract

def _ocr_pdf(data: bytes) -> str:
    try:
        import fitz
        from PIL import Image
    except ImportError as exc:
        raise RuntimeError("PDF OCR needs PyMuPDF and Pillow.") from exc
    pytesseract = _ocr_engine()
    pages = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        matrix = fitz.Matrix(_OCR_DPI / 72, _OCR_DPI / 72)
        for page in doc:
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            pages.append(pytesseract.image_to_string(img))
    return "\n".join(pages)

def _ocr_image(data: bytes) -> str:
    from PIL import Image
    pytesseract = _ocr_engine()
    return pytesseract.image_to_string(Image.open(io.BytesIO(data)))
